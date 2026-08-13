#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
投稿版 docx 自动验证脚本 v0.1.0（journal-submission-docx 技能）
生成后必须运行；任一 FAIL 即需修复后再交付。

用法：
  python3 validate_submission.py <投稿版.docx> [<正文md目录>]

检查项（均为实测踩坑沉淀）：
  1. HTML 标签残留（<sup>/<sub>/<i>/<b> 等）——段落与表格单元格
  2. 三线表：仅 top/bottom 粗线 + 表头下栏目线；insideH/insideV/left/right 全 none
  3. 列宽生效：tblGrid 存在且每列宽 > 0（只设 cell.width 会失效）
  4. 上标星号：表格中含 * 的单元格应有 superscript run（无上标则报警）
  5. 图片完整性：md 引用图数 == docx 嵌入图数（缺图报警）
  6. 页码域：footer 含 PAGE 域
  7. 文件属性：author 为空（匿名要求）
"""
import re, os, sys
from docx import Document
from docx.oxml.ns import qn

def check_html_residue(doc):
    bad = []
    for p in doc.paragraphs:
        if re.search(r"</?[a-zA-Z]+>", p.text):
            bad.append(p.text[:50])
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if re.search(r"</?[a-zA-Z]+>", c.text):
                    bad.append(c.text[:50])
    return bad

def check_three_line_table(doc):
    problems = []
    for ti, t in enumerate(doc.tables):
        tblPr = t._tbl.tblPr
        borders = tblPr.find(qn("w:tblBorders"))
        if borders is None:
            problems.append(f"表{ti+1}: 无 tblBorders"); continue
        d = {c.tag.split("}")[1]: c.get(qn("w:val")) for c in borders}
        if d.get("top") != "single" or d.get("bottom") != "single":
            problems.append(f"表{ti+1}: 顶线/底线缺失 top={d.get('top')} bottom={d.get('bottom')}")
        for edge in ("insideH", "insideV", "left", "right"):
            if d.get(edge) != "none":
                problems.append(f"表{ti+1}: {edge}={d.get(edge)}（应为 none）")
        # 表头下栏目线（跨列表头时在第 1 行，普通表头在第 0 行；前两行任一单元格有 tcBorders 即通过）
        has_hdr_line = False
        for ri in range(min(2, len(t.rows))):
            for ci in range(len(t.columns)):
                tc = t.cell(ri, ci)._tc
                tcPr = tc.find(qn("w:tcPr"))
                if tcPr is not None and tcPr.find(qn("w:tcBorders")) is not None:
                    has_hdr_line = True
                    break
            if has_hdr_line:
                break
        if not has_hdr_line:
            problems.append(f"表{ti+1}: 表头无栏目线（tcBorders bottom）")
    return problems

def check_col_width(doc):
    problems = []
    for ti, t in enumerate(doc.tables):
        grid = t._tbl.find(qn("w:tblGrid"))
        if grid is None:
            problems.append(f"表{ti+1}: 无 tblGrid（列宽未生效）"); continue
        cols = grid.findall(qn("w:gridCol"))
        for j, gc in enumerate(cols):
            w = gc.get(qn("w:w"))
            if not w or int(w) <= 0:
                problems.append(f"表{ti+1} 列{j+1}: 列宽为 {w}（列宽未生效）")
    return problems

def check_superscript(doc):
    """含星号的表格单元格应有上标 run；无星号不要求"""
    missing = []
    for ti, t in enumerate(doc.tables):
        for ri, r in enumerate(t.rows):
            for ci, c in enumerate(r.cells):
                if "*" in c.text:
                    has_sup = any(run.font.superscript for p in c.paragraphs for run in p.runs)
                    if not has_sup:
                        missing.append(f"表{ti+1} 行{ri+1} 列{ci+1}: 含星号但无上标 [{c.text[:25]}]")
    return missing

def check_images(doc, md_dir=None):
    n_docx = len(doc.inline_shapes)
    if md_dir is None:
        return n_docx, None
    n_md = 0
    missing_files = []
    for f in sorted(os.listdir(md_dir)):
        if not f.endswith(".md") or not re.match(r"^\d\d-", f):
            continue
        text = open(os.path.join(md_dir, f), encoding="utf-8").read()
        for m in re.finditer(r"!\[([^\]]*)\]\(([^)]+)\)", text):
            n_md += 1
            if not os.path.exists(os.path.join(md_dir, m.group(2))):
                missing_files.append(f"{f}: {m.group(2)}")
    return n_docx, (n_md, missing_files)

def check_page_number(doc):
    for sec in doc.sections:
        for p in sec.footer.paragraphs:
            if "PAGE" in p._p.xml:
                return True
    return False

def check_author(doc):
    return (doc.core_properties.author or "").strip() == ""

def main():
    docx_path, md_dir = sys.argv[1], (sys.argv[2] if len(sys.argv) > 2 else None)
    doc = Document(docx_path)
    fails, warns = [], []

    r = check_html_residue(doc)
    if r: fails.append(f"HTML 标签残留 {len(r)} 处: {r[:3]}")
    else: print("✓ HTML 标签残留: 0")

    r = check_three_line_table(doc)
    if r: fails.append("三线表异常:\n  " + "\n  ".join(r[:5]))
    else: print(f"✓ 三线表: {len(doc.tables)} 表全部仅三线")

    r = check_col_width(doc)
    if r: fails.append("列宽异常:\n  " + "\n  ".join(r[:5]))
    else: print("✓ 列宽: 全部生效（tblGrid）")

    r = check_superscript(doc)
    if r:
        warns.append(f"星号无上标 {len(r)} 处: {r[:3]}")
    else: print("✓ 上标星号: 全部上标")

    n_docx, md_info = check_images(doc, md_dir)
    if md_info is not None:
        n_md, missing = md_info
        if missing: fails.append(f"md 引用的图片文件缺失: {missing}")
        if n_docx != n_md:
            fails.append(f"图片数不符: docx={n_docx}, md={n_md}")
        else: print(f"✓ 图片: {n_docx}/{n_md} 全部嵌入")
    else:
        print(f"✓ 图片: {n_docx} 张嵌入（未提供 md 目录比对）")

    if check_page_number(doc): print("✓ 页码域: 存在")
    else: warns.append("页码域缺失")

    if check_author(doc): print("✓ 文件属性: 已匿名")
    else: warns.append("文件属性 author 非空（需清理）")

    print()
    if fails:
        print("❌ FAIL（必须修复）:")
        for f in fails: print("  -", f)
        sys.exit(1)
    if warns:
        print("⚠️  WARN（建议处理）:")
        for w in warns: print("  -", w)
    print("✅ 验证通过（警告项见上）")

if __name__ == "__main__":
    main()
