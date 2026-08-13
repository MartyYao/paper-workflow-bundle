#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
期刊投稿版 docx 生成器 v0.1.0（journal-submission-docx 技能）
正式路径：python-docx 逐元素渲染，期刊样式配置驱动。

用法：
  python3 generate_submission.py <期刊配置名> <正文md目录> <输出目录>

当前配置：中国工业经济（2026-08-11 实测版）
"""
import re, os, sys, subprocess, tempfile, zipfile
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============ 期刊配置 ============
# size: pt；line: 行距倍数；indent: 缩进（字符数）；align: left/center/justify
CFG = {
    "中国工业经济": {
        "journal_name": "中国工业经济",
        "page": {"width": Cm(21.0), "height": Cm(29.7),
                 "top": Cm(2.54), "bottom": Cm(2.54), "left": Cm(3.17), "right": Cm(3.17)},
        "title":  {"font": "黑体", "size": 16, "bold": True, "align": "center", "line": 1.5},
        "abstract_label": {"font": "黑体", "size": 10.5, "bold": False},
        "abstract": {"font": "宋体", "size": 10.5, "line": 1.5, "first_indent": 2},
        "h1": {"font": "仿宋", "size": 14, "bold": False, "align": "center", "line": 1.5},   # 一、四号仿宋居中
        "h2": {"font": "黑体", "size": 10.5, "bold": False, "indent": 2, "line": 1.5},        # 1.五号黑体缩2字
        "h3": {"font": "宋体", "size": 10.5, "bold": False, "indent": 2, "line": 1.5},        # （1）五号宋体
        "body": {"font": "宋体", "size": 10.5, "line": 1.5, "first_indent": 2},               # 正文五号宋体
        "table_caption": {"font": "黑体", "size": 10.5, "align": "center", "line": 1.5},      # 表名在上
        "table_note": {"font": "宋体", "size": 7.5, "line": 1.0},                               # 表注：六号、单倍行距（比正文小）
        "fig_caption": {"font": "黑体", "size": 10.5, "align": "center", "line": 1.5},        # 图名在下
        "ref": {"font": "宋体", "size": 10.5, "line": 1.5},
    },
}

ALIGN = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
         "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}

def set_run(run, font, size, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)

def latex_to_omml(latex, inline=False, size_pt=10.5):
    """LaTeX → OMML（WPS 兼容公式，2026-08-11 最终定案）。
    链路：latex2mathml（LaTeX→MathML）→ mathml2omml（MathML→OMML）→ set_formula_font（WPS 格式）。
    WPS 渲染 OMML 的关键：字体写在 w:rPr（DejaVu Math TeX Gyre）——复刻 WPS 转换器的输出格式。
    size_pt：公式字号跟随上下文（正文五号/表注六号/表格小字号），避免突兀"""
    from latex2mathml.converter import convert as l2m
    import mathml2omml
    mml = l2m(latex)
    omml = mathml2omml.convert(mml)
    omml = set_formula_font(omml, size_pt=size_pt)
    return [omml]

def set_formula_font(omml, size_pt=10.5):
    """WPS 兼容公式格式（2026-08-11 定案，分析 WPS 转换后的公式 XML 复刻）：
    m:r 内子元素顺序 = m:rPr（空）→ w:rPr（rFonts DejaVu Math TeX Gyre + sz）→ m:t。
    字体必须写在 w:rPr（非 m:rPr）——WPS 只认 w:rPr 里的 rFonts；
    字体 DejaVu Math TeX Gyre 是 WPS 自带数学字体（非 Cambria/STIX）。
    sz = size_pt*2（半磅），字号跟随上下文。"""
    from lxml import etree
    ns_m = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    # mathml2omml 输出无命名空间声明，解析前补上
    omml = omml.replace("<m:oMath>", f'<m:oMath xmlns:m="{ns_m}" xmlns:w="{ns_w}">', 1)
    root = etree.fromstring(omml)
    for r in root.findall(f".//{{{ns_m}}}r"):
        # 1) 删除 m:rPr 里的 w:rFonts（非标准位置，WPS 不认）
        rPr = r.find(f"{{{ns_m}}}rPr")
        if rPr is not None:
            for rf in rPr.findall(f"{{{ns_w}}}rFonts"):
                rPr.remove(rf)
        # 2) 删除已有 w:rPr（避免重复）
        for old in r.findall(f"{{{ns_w}}}rPr"):
            r.remove(old)
        # 3) 在 m:t 之前插入 w:rPr（rFonts DejaVu Math TeX Gyre + sz 五号 21）
        t = r.find(f"{{{ns_m}}}t")
        if t is not None:
            w_rPr = etree.Element(f"{{{ns_w}}}rPr")
            rf = etree.SubElement(w_rPr, f"{{{ns_w}}}rFonts")
            rf.set(f"{{{ns_w}}}hint", "default")
            rf.set(f"{{{ns_w}}}ascii", "DejaVu Math TeX Gyre")
            rf.set(f"{{{ns_w}}}hAnsi", "DejaVu Math TeX Gyre")
            rf.set(f"{{{ns_w}}}eastAsia", "宋体")
            sz = etree.SubElement(w_rPr, f"{{{ns_w}}}sz")
            sz.set(f"{{{ns_w}}}val", str(int(size_pt * 2)))
            r.insert(list(r).index(t), w_rPr)
    return etree.tostring(root, encoding="unicode")

def add_omml_into_paragraph(p, omml_xml):
    """把 OMML 公式 XML 插入段落（python-docx oxml）"""
    from docx.oxml import parse_xml
    el = parse_xml(omml_xml)
    p._p.append(el)

def mathify_text(text):
    """正文/表格中的数学符号文本 → 行内 LaTeX（$..$），由 add_rich_runs 转 OMML。
    解决正文公式符号未排版问题（dev_it、|ε|、max(0,−ε)、ln(1+Subsidy_it)、R²）。
    白名单模式，避免误伤普通文本；已有 $..$ 公式段受保护不重复处理。"""
    parts = re.split(r"(\$[^$\n]+\$)", text)
    out = []
    for seg in parts:
        if seg.startswith("$") and seg.endswith("$"):
            out.append(seg)   # 已是 LaTeX 公式，不处理
            continue
        t = re.sub(r'ln\(1\+([A-Za-z]+_[a-z]+)\)', lambda m: r'$\ln(1+%s)$' % m.group(1), seg)
        t = re.sub(r'ln\(([A-Za-z]+)_([a-z]{1,3})-?(\d*)\)', lambda m: r'$\ln(%s_{%s%s})$' % (m.group(1), m.group(2), ('-' + m.group(3)) if m.group(3) else ''), t)
        t = re.sub(r'\|ε\|=under\+over', lambda m: r'$|\varepsilon|=under+over$', t)
        t = re.sub(r'max\(0,\s*−?ε\)', lambda m: r'$\max(0,-\varepsilon)$', t)
        t = re.sub(r'\|ε\|', lambda m: r'$|\varepsilon|$', t)
        t = re.sub(r'R²', lambda m: r'$R^{2}$', t)
        # 希腊字母下标（μ_i、λ_t、ε_it、μ_ijt → \mu_{i} 等）
        GREEK = {'μ': r'\mu', 'λ': r'\lambda', 'ε': r'\varepsilon', 'α': r'\alpha',
                 'β': r'\beta', 'γ': r'\gamma', 'δ': r'\delta', 'ρ': r'\rho', 'σ': r'\sigma'}
        t = re.sub(r'([μλεαβγδρσ])_([a-z]{1,3})\b', lambda m: r'$%s_{%s}$' % (GREEK[m.group(1)], m.group(2)), t)
        # ASCII 希腊名（mu_i/lambda_t/beta_k → \mu_{i} 等，须在通用下标模式前）
        t = re.sub(r'\b(mu|lambda|beta|alpha|gamma|delta|rho|sigma)_([a-z]{1,3})\b',
                   lambda m: r'$\\%s_{%s}$' % (m.group(1), m.group(2)), t)
        # 变量下标模式（dev_it、Subsidy_it、X_it 等，单字母变量也支持）
        t = re.sub(r'\b([A-Za-z]{1,})_([a-z]{1,3})\b', lambda m: r'$%s_{%s}$' % (m.group(1), m.group(2)), t)
        out.append(t)
    return "".join(out)

def add_rich_runs(p, text, cfg, bold=False, italic=False):
    """富文本渲染：<sup>..</sup> 上标、**..** 加粗、*..* 斜体、$..$ 行内公式（OMML）；
    数学符号文本（dev_it/|ε|/max(0,−ε) 等）自动转行内公式"""
    text = mathify_text(text)
    if any(ord(c) < 32 and c not in '\n\t' for c in text):
        print(f'[控制字符] 段落文本: {repr(text[:200])}')
    for seg in re.split(r"(\$[^$\n]+\$|<sup>.*?</sup>|\*\*.*?\*\*|\*.*?\*)", text):
        if not seg:
            continue
        if seg.startswith("$") and seg.endswith("$"):
            # 行内公式 → WPS 兼容 OMML（字号跟随上下文）
            for omml in latex_to_omml(seg[1:-1], inline=True, size_pt=cfg.get("size", 10.5)):
                add_omml_into_paragraph(p, omml)
            continue
        if seg.startswith("<sup>"):
            run = p.add_run(seg[5:-6])
            set_run(run, cfg["font"], cfg["size"], bold=bold, italic=italic)
            run.font.superscript = True
        elif seg.startswith("**"):
            run = p.add_run(seg[2:-2])
            set_run(run, cfg["font"], cfg["size"], bold=True, italic=italic)
        elif seg.startswith("*"):
            run = p.add_run(seg[1:-1])
            set_run(run, cfg["font"], cfg["size"], bold=bold, italic=True)
        else:
            run = p.add_run(seg)
            set_run(run, cfg["font"], cfg["size"], bold=bold, italic=italic)

def add_para(doc, text, style, bold=False, italic=False, keep_with_next=False, **kw):
    """按配置添加段落；支持富文本标记；段落间距按样式收紧（表上下间距自适应）"""
    cfg = CFG["中国工业经济"][style]
    # 先创建隔离空段（h1 时）——必须在标题段之前，否则起不到隔离作用
    if style == "h1" and text.strip() != "摘要":
        # 章节另起一页：**隔离空段 + 标题段 page_break_before**。
        # 坑：①page_break_before 在"表格/表注/OMML 公式后紧跟"的段落被 LO 吞（标题消失）；
        # ②w:br type=page 放空段也不渲染分页。解法：先插普通空段隔离特殊内容，
        # 再给标题段设 page_break_before（此时标题前是普通段，LO 渲染正常）
        # 摘要例外：紧跟论文标题，同页不另起
        iso = doc.add_paragraph()
        isopf = iso.paragraph_format
        isopf.line_spacing = 1.0
        isopf.space_before = Pt(0)
        isopf.space_after = Pt(0)
        iso.add_run("")
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = ALIGN.get(cfg.get("align", "left"), WD_ALIGN_PARAGRAPH.LEFT)
    pf.line_spacing = cfg.get("line", 1.5)
    SPACING = {"table_caption": (6, 2), "table_note": (0, 6), "fig_caption": (2, 6), "title": (0, 6)}
    sb, sa = SPACING.get(style, (0, 0))
    pf.space_before = Pt(sb)
    pf.space_after = Pt(sa)
    extra_sb = kw.pop("space_before_extra", 0)
    if extra_sb:
        pf.space_before = Pt(pf.space_before.pt + extra_sb) if pf.space_before else Pt(extra_sb)
    if style == "h1" and text.strip() != "摘要":
        pf.page_break_before = True
    if keep_with_next:
        pf.keep_with_next = True
    if "first_indent" in cfg:
        pf.first_line_indent = Pt(cfg["size"] * cfg["first_indent"])
    if "indent" in cfg:
        pf.left_indent = Pt(cfg["size"] * cfg["indent"])
    add_rich_runs(p, text, cfg, bold=bold, italic=italic)
    return p

def is_group_row(row):
    """分组行：除第 1 列外其他列全空（如 'ln 审计费用：总样本'）"""
    return all(not c.strip() for c in row[1:])

def split_header(rows):
    """表头拆分：'前缀：后缀' 模式 → 前缀行（相邻同前缀合并）+ 后缀行。
    例：['变量','ln 审计费用：总样本','ln 审计费用：不足组',...]
    → 前缀行：['变量','ln 审计费用'(跨3列),'ln 专利引用'(跨3列)]
      后缀行：['变量','总样本','不足组','超发组','总样本','不足组','超发组']
    返回 (prefixes, suffixes) 或 None（无 "：" 模式时）"""
    hdr = rows[0]
    prefixes, suffixes = [], []
    for c in hdr:
        t = strip_md(c)
        found = False
        for sep in ("：", ":"):
            if sep in t:
                p, s = t.split(sep, 1)
                prefixes.append(p.strip())
                suffixes.append(s.strip())
                found = True
                break
        if not found:
            prefixes.append(None)
            suffixes.append(c.strip())
    if not any(p is not None for p in prefixes):
        return None
    return prefixes, suffixes

def round_cell(text, decimals=3):
    """表格数字四舍五入到指定位数（默认三位；宽表 >12 列用两位）。
    保留 <sup> 星号标记（相关系数表等）；只处理纯数字单元格，其他文本不动。"""
    m = re.match(r"^(-?\d+\.\d{4,})(<sup>.*?</sup>)?$", text)
    if m:
        return f"{round(float(m.group(1)), decimals):.{decimals}f}{m.group(2) or ''}"
    m = re.match(r"^\((-?\d+\.\d{4,})\)$", text)
    if m:
        return f"({round(float(m.group(1)), decimals):.{decimals}f})"
    return text

def set_cell(cell, text, size, bold=False, align_left=False, decimals=3):
    """单元格填充：垂直居中、单倍行距、段间距 0（行高自适应收紧）；
    数字单元格统一三位小数（宽表两位，round_cell），表头/文字不受影响"""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if align_left else WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    add_rich_runs(p, round_cell(text, decimals), {"font": "宋体", "size": size}, bold=bold)

def add_table(doc, rows, note=None):
    """三线表：仅顶线/底线 1.5pt + 栏目线 0.75pt，无竖线、无中间横线
    列宽按内容自适应（分组行/跨列表头不参与估算），长内容自动降字号；
    分组行跨列合并加粗；表头 '前缀：后缀' 自动拆分为跨列合并的两行表头；
    note：表注并入表格最后一行（跨列合并、无边框）——表注与表格不可拆分"""
    ncol = len(rows[0])
    header = split_header(rows)
    n_extra = 1 if header else 0          # 跨列表头时表头占 2 行
    data_rows = [r for r in rows if not is_group_row(r)] or rows
    if header:
        data_rows = rows[1:]
    total_rows = len(rows) + n_extra + (1 if note else 0)
    table = doc.add_table(rows=total_rows, cols=ncol)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT   # 与正文左对齐（不居中，铺满版心）
    table.autofit = False
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge, sz in (("top", 12), ("bottom", 12)):   # 仅顶线、底线
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), str(sz)); el.set(qn("w:color"), "000000")
        borders.append(el)
    for edge in ("left", "right", "insideH", "insideV"):  # 其余全 none
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    tblPr.append(borders)
    # 栏目线：suffix 行（跨列表头时第 1 行，否则表头第 0 行）下边框 0.75pt
    hdr_row_idx = 1 if header else 0
    for j in range(ncol):
        tc = table.cell(hdr_row_idx, j)._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6"); bottom.set(qn("w:color"), "000000")
        tcBorders.append(bottom)
        tcPr.append(tcBorders)
    # 列宽布局：按内容需求分配（数据行+后缀表头参与估算），超宽逐档降字号，极限等比压缩
    prefixes = suffixes = None
    if header:
        prefixes, suffixes = header
        est_rows = [list(suffixes)] + data_rows
    else:
        est_rows = data_rows
    size, widths, decimals, margin = calc_table_layout(est_rows)
    hsize = min(size, 8.5)   # 表头字号：cap 8.5pt（用户 2026-08-11：表头字号过大导致换行，8.5pt 使窄列表头单行）
    # 总宽铺满版心：按内容需求比例放大，保证与正文等宽（LEFT 不居中）
    scale = 14.6 / sum(widths)
    widths = [w * scale for w in widths]
    # 单元格边距按表宽收紧（默认 0.19cm；>12 列宽表 0.1cm），为数字单行挤出空间
    tblPr2 = table._tbl.tblPr
    cellMar = OxmlElement("w:tblCellMar")
    m_twips = int(margin * 567)
    for edge, v in (("left", m_twips), ("right", m_twips), ("top", 28), ("bottom", 28)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(v))
        el.set(qn("w:type"), "dxa")
        cellMar.append(el)
    tblPr2.append(cellMar)
    for j in range(ncol):
        table.columns[j].width = Cm(widths[j])   # 更新 tblGrid（关键：不设则列宽不生效）

    # 渲染：跨列表头（前缀行合并 + 后缀行）→ 数据行/分组行
    if header:
        j = 0
        while j < ncol:
            if prefixes[j] is None:
                j += 1
                continue
            # 合并相邻同前缀列（while 推进：for 循环变量被改不生效，会重复处理已合并列）
            end = j
            while end + 1 < ncol and prefixes[end + 1] == prefixes[j]:
                end += 1
            if end > j:
                merged = table.cell(0, j)
                for k in range(j + 1, end + 1):
                    merged = merged.merge(table.cell(0, k))
                set_cell(merged, prefixes[j], hsize, bold=True)
            else:
                set_cell(table.cell(0, j), prefixes[j], hsize, bold=True)
            j = end + 1
        for j in range(ncol):
            set_cell(table.cell(1, j), suffixes[j], hsize, bold=True, align_left=(j == 0))

    for i, row in enumerate(rows[1:] if header else rows):
        ti = i + n_extra + (1 if header else 0)   # header 占 prefix+suffix 两行
        tr = table.rows[ti]._tr
        trPr = tr.get_or_add_trPr()
        trPr.append(OxmlElement("w:cantSplit"))   # 行内不分页
        if ti <= hdr_row_idx:
            trPr.append(OxmlElement("w:tblHeader"))  # 表头跨页重复
        if is_group_row(row):
            # 分组行：跨列合并 + 加粗左对齐
            merged = table.cell(ti, 0)
            for j in range(1, ncol):
                merged = merged.merge(table.cell(ti, j))
            set_cell(merged, row[0].strip(), size, bold=True, align_left=True)
            continue
        for j, cell_text in enumerate(row):
            cell = table.cell(ti, j)
            cell.width = Cm(widths[j])
            set_cell(cell, cell_text.strip(), size, bold=(ti == hdr_row_idx), align_left=(j == 0), decimals=decimals)
    # 表注行：并入表格（跨列合并、四边无边框、六号左对齐）——与表格不可拆分
    if note:
        ni = total_rows - 1
        note_cell = table.cell(ni, 0)
        for j in range(1, ncol):
            note_cell = note_cell.merge(table.cell(ni, j))
        tcPr = note_cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for edge in ("top", "bottom", "left", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "none")
            tcBorders.append(el)
        tcPr.append(tcBorders)
        note_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = note_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.line_spacing = 1.0
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        add_rich_runs(p, note, {"font": "宋体", "size": 7.5})   # 表注六号
    return None

def strip_md(text):
    """去除 md/HTML 标记（<sup>..</sup>、**、*），只留显示文本"""
    return re.sub(r"<sup>.*?</sup>|\*\*|\*", "", text)

def text_width_cm(text, size):
    """估算文本渲染宽度：CJK 全宽，英文/数字半宽（按去标签后的显示文本）"""
    w = 0.0
    for ch in strip_md(text):
        w += 1.0 if ord(ch) > 0x2E80 else 0.5
    return w * size * 0.037

def col_need(j, rows, size, decimals=3, margin=0.19):
    """单列需求：短列（≤14 显示字符）按单行宽度估算；长文本列（定义类）允许换行，只占下限 3.2cm
    表头长文本（如 'ln 审计费用：总样本'）允许两行换行，需求封顶 1.8cm；
    数字按指定位数后的显示宽度估算（round_cell），列宽含单元格边距 margin×2"""
    texts = [strip_md(round_cell(r[j], decimals)) for r in rows]
    maxlen = max(len(t) for t in texts)
    if maxlen <= 14:
        need = max(text_width_cm(t, size) for t in texts) + margin * 2 + 0.05
        hdr = strip_md(rows[0][j])
        if len(hdr) > 8:
            need = min(need, 1.8)
        return need
    return 3.2

def calc_table_layout(rows, avail_cm=14.6):
    """列宽布局：短列按单行需求、长列按换行下限；超宽逐档降字号（五号→小五→六号→6.5→6）。
    宽表（>12 列，如相关系数矩阵）自动降精度至 2 位小数 + 边距收紧 0.1cm（15 列纵向单行）；
    纯纵向方案（无分节）：LibreOffice/WPS/Word 渲染一致"""
    ncol = len(rows[0])
    decimals = 2 if ncol > 12 else 3
    margin = 0.1 if ncol > 12 else 0.19
    for size in (10.5, 9.0, 7.5, 6.5, 6.0):
        needs = [col_need(j, rows, size, decimals, margin) for j in range(ncol)]
        if sum(needs) <= avail_cm:
            return size, needs, decimals, margin
    # 6pt 仍超（极少）：等比压缩，保证总宽 = 可用宽度
    size = 6.0
    needs = [col_need(j, rows, size, decimals, margin) for j in range(ncol)]
    scale = avail_cm / sum(needs)
    return size, [n * scale for n in needs], decimals, margin

def insert_section_break(p, landscape):
    """在段落 p 的 pPr 中嵌入连续分节符（w:type=continuous，不强制分页）。
    LibreOffice/WPS 语义：该段结束当前节，sectPr 描述刚结束的节。"""
    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement("w:sectPr")
    stype = OxmlElement("w:type")
    stype.set(qn("w:val"), "continuous")
    sectPr.append(stype)
    pgSz = OxmlElement("w:pgSz")
    if landscape:
        pgSz.set(qn("w:w"), str(int(29.7 * 567)))
        pgSz.set(qn("w:h"), str(int(21.0 * 567)))
        pgSz.set(qn("w:orient"), "landscape")
    else:
        pgSz.set(qn("w:w"), str(int(21.0 * 567)))
        pgSz.set(qn("w:h"), str(int(29.7 * 567)))
    sectPr.append(pgSz)
    pgMar = OxmlElement("w:pgMar")
    for attr, v in (("top", 1440), ("bottom", 1440), ("left", 1800), ("right", 1800),
                    ("header", 851), ("footer", 992), ("gutter", 0)):
        pgMar.set(qn(f"w:{attr}"), str(v))
    sectPr.append(pgMar)
    pPr.append(sectPr)

def add_figure(doc, img_path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True   # 图注跟随图片，同页不拆分
    p.add_run().add_picture(img_path, width=Cm(14.0))
    add_para(doc, caption, "fig_caption")

def add_page_number(doc):
    """页脚居中插入 PAGE 页码域"""
    sec = doc.sections[0]
    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)

def parse_and_build(md_dir, out_docx):
    doc = Document()
    sec = doc.sections[0]
    pg = CFG["中国工业经济"]["page"]
    sec.page_width, sec.page_height = pg["width"], pg["height"]
    sec.top_margin, sec.bottom_margin = pg["top"], pg["bottom"]
    sec.left_margin, sec.right_margin = pg["left"], pg["right"]
    add_page_number(doc)

    files = sorted(f for f in os.listdir(md_dir) if f.endswith(".md") and not f.startswith("."))
    files = [f for f in files if re.match(r"^\d\d-", f)]  # 仅正文分节文件
    ref_section = False
    pending_caption = None      # 表格名（表上方）
    pending_rows = []           # 表格行收集
    pending_note = None         # 表注（并入表格最后一行）
    after_table = False         # 表格后状态（表注识别）
    title_done = False          # 论文标题已渲染（00-摘要.md 第一行 # 题目）
    formula_open = False        # 块公式等待编号行（（N））

    last_was_table = [False]

    def para(text, style, **kw):
        """add_para 包装：纯纵向方案；表格后紧跟的段落（正文/标题）加 6pt 上间距，
        避免正文/标题贴住表注（用户 2026-08-11）"""
        if last_was_table[0]:
            kw["space_before_extra"] = 6
            last_was_table[0] = False
        return add_para(doc, text, style, **kw)

    def flush_table():
        nonlocal pending_caption, pending_rows, pending_note, after_table
        if pending_rows:
            if pending_caption:
                add_para(doc, pending_caption, "table_caption", keep_with_next=True)
            add_table(doc, pending_rows, note=pending_note)
            pending_caption, pending_rows, pending_note, after_table = None, [], None, True
            last_was_table[0] = True

    for fn in files:
        path = os.path.join(md_dir, fn)
        with open(path, encoding="utf-8") as f:
            lines = f.read().split("\n")
        for raw in lines:
            line = raw.rstrip()
            if not line.strip():
                continue   # 空行仅作分隔，不触发 flush（否则表格后的"注："会被错配）
            # 状态行/文件标题剔除；00-摘要.md 第一行 # 为论文标题（渲染为标题样式）
            if line.startswith("> 状态"):
                continue
            if re.match(r"^# [^#]", line):
                if not title_done and fn.startswith("00-"):
                    para(line[2:].strip(), "title")
                    title_done = True
                continue
            # 块公式（$$..$$ 单行）→ WPS 兼容 OMML 居中；下一行（N）为编号（右对齐）
            if line.startswith("$$") and line.endswith("$$"):
                flush_table()
                latex = line[2:-2].strip()
                fp = doc.add_paragraph()
                fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                fpf = fp.paragraph_format
                fpf.line_spacing = 1.5
                for omml in latex_to_omml(latex, size_pt=10.5):
                    add_omml_into_paragraph(fp, omml)
                formula_open = True
                continue
            if formula_open and re.match(r"^（\d+）$", line.strip()):
                # 公式编号：右对齐制表位
                fp.paragraph_format.tab_stops.add_tab_stop(Cm(14.6), WD_TAB_ALIGNMENT.RIGHT)
                run = fp.add_run("\t" + line.strip())
                set_run(run, "宋体", 10.5)
                formula_open = False
                continue
            if line.startswith("## "):
                flush_table()
                txt = line[3:].strip()
                if "参考文献" in txt:
                    ref_section = True
                para(txt, "h1")
                continue
            # markdown 分隔线（---）跳过
            if line.strip() == "---":
                continue
            if line.startswith("### "):
                flush_table()
                para(line[4:].strip(), "h2")
                continue
            if line.startswith("#### "):
                flush_table()
                para(line[5:].strip(), "h3")
                continue
            # 表名（**表 N xxx** 独立段）
            m = re.match(r"^\*\*表\s*([^*]+?)\*\*$", line)
            if m:
                flush_table()
                pending_caption = f"表{m.group(1).strip()}"
                continue
            # 表格行
            if line.startswith("|"):
                cells = [c.strip() for c in line.strip("|").split("|")]
                if re.match(r"^:?-+:?$", cells[0].replace(" ", "")):
                    continue  # 分隔行
                pending_rows.append(cells)
                after_table = True   # 上一行是表格行（表格后紧邻段可能是表注）
                continue
            # 图片
            m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line)
            if m:
                flush_table()
                cap = m.group(1).strip()
                img_rel = m.group(2)
                img_path = os.path.join(md_dir, img_rel)
                if os.path.exists(img_path):
                    add_figure(doc, img_path, cap)
                else:
                    para(f"【缺图】{cap}（{img_rel}）", "body")
                continue
            # 表注（表格后紧邻的"注："段）——表格未渲染时并入表格最后一行（不可拆分）
            if after_table and line.startswith(("注：", "资料来源：")):
                if pending_rows:
                    pending_note = line
                else:
                    para(line, "table_note")   # 注到达时表格已渲染（罕见）：独立表注段
                after_table = False   # 已消费，防止后续段误判
                continue
            flush_table()
            # 参考文献条目
            if ref_section and not line.startswith("#"):
                para(line, "ref")
            else:
                para(line, "body")   # 含 $..$ 行内公式时由 add_rich_runs 转 OMML
    flush_table()

    # 匿名要求：清空文件属性（生成器内集成，避免遗漏）
    cp = doc.core_properties
    cp.author = ""
    cp.last_modified_by = ""
    cp.comments = ""

    os.makedirs(os.path.dirname(out_docx), exist_ok=True)
    doc.save(out_docx)
    return out_docx

if __name__ == "__main__":
    md_dir, out_arg = sys.argv[1], sys.argv[2]
    # 命名规则（用户 2026-08-11）：用什么模板生成 → "投稿版-<期刊名>.docx"
    if out_arg.endswith(".docx"):
        out_docx = out_arg   # 显式路径优先（兼容旧调用）
    else:
        journal = next(iter(CFG))   # 当前配置期刊
        out_docx = os.path.join(out_arg, f"投稿版-{journal}.docx")
    print("输出:", parse_and_build(md_dir, out_docx))
